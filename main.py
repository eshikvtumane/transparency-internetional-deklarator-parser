# -*- encoding: utf-8 -*-
import docx

from auto import Auto
from deklarator import DeklaratorXML


class Depute(object):
    def __init__(self):
        self.name = None
        self.position = None
        self.salary = None
        self.objects_owner = []
        self.objects_not_owner = []
        self.machines = []
        self._machines_list = []

    def set_name(self, name):
        self.name = name

    def set_salary(self, salary):
        salary = self.clear_number(salary)

        if salary != '':
            self.salary = float(salary)

    def set_object_owner(self, object_type, object_size, country):
        object_list = [object_type, self.clear_number(object_size), country]
        if object_list != ['', '', ''] and object_list not in self.objects_owner:
            self.objects_owner.append(object_list)

    def set_object_not_owner(self, object_type, object_size, country):
        object_list = [object_type, self.clear_number(object_size), country]
        if object_list != ['', '', ''] and object_list not in self.objects_not_owner:
            self.objects_not_owner.append(object_list)

    def set_machine(self, machine):
        if machine != '' and machine not in self._machines_list:
            self._machines_list.append(machine)
            type_machine = ''
            machines = []

            for machine in machine.split(';'):
                machines += machine.split('\n')

            for machine in machines:
                machine = machine.strip()
                if ':' in machine:
                    type_machine = machine.split(':')[0]

                    if 'автомобили легковые' in type_machine or 'автомобиль легковой' in type_machine or 'легковые автомобили' in type_machine:
                        type_machine = 'а/м легковой'
                    elif 'автомобили грузовые' in type_machine or 'грузовые автомобили' in type_machine or 'грузовой автомобиль' in type_machine:
                        type_machine = 'а/м грузовой'
                    elif 'моторные лодки' in type_machine:
                        type_machine = 'Моторная лодка'
                    elif 'иные транспортные средства' in type_machine:
                        type_machine = type_machine.replace('иные транспортные средства ', '')
                    elif 'водный транспорт' in type_machine:
                        type_machine = type_machine.replace('водный транспорт', '')
                    elif 'водный мотоцикл' in type_machine:
                        type_machine = type_machine.replace('водный мотоцикл', '')
                    elif 'мототранспортные средства' in type_machine:
                        type_machine = type_machine.replace('мототранспортное средство', '')
                    elif 'водный мотоцикл' in type_machine:
                        type_machine = type_machine.replace('водный мотоцикл', '')
                    elif Auto.is_auto(type_machine):
                        machine = type_machine
                        type_machine = ''
                    else:
                        type_machine = type_machine
                    machine = machine.split(':')[1].strip()

                    if machine:
                        self.machines.append(('%s %s' % (type_machine, machine)).strip().replace('  ', ' '))
                # elif '' == machine:
                #     type_machine = ''
                else:
                    if machine != '':
                        machine_dict = [
                            'водный транспорт', 'автомобили легковые', 'автомобиль легковой', 'легковой автомобиль',
                            'мототранспортное средство мотоцикл', 'мототранспортное средство',
                            'бульдозер', 'иные транспортные средства', 'автомобиль грузовой', 'грузовые автомобили',
                            'водный мотоцикл', 'грузовой автомобиль', 'грузовой автомобиль', 'водный мотоцикл'
                        ]

                        if machine in machine_dict:
                            type_machine = machine
                            continue

                        elif Auto.is_auto(type_machine):
                            machine = type_machine
                            type_machine = ''

                        self.machines.append(('%s %s' % (type_machine, machine)).strip().replace('  ', ' '))
                    # self.machines.append(machine.strip())

    def clear_number(self, number):
        return number.replace(' ', '').replace(' ', '').replace(',', '.')

    def write_owner_for_old_files(self, owner, square, country):
        if 'собственность' in owner or 'совместная собственность' in owner \
                or 'долевая собственность' in owner or 'доли' in owner:
            self.set_object_owner(owner, square, country)
        else:
            self.set_object_not_owner(owner, square, country)

    def __str__(self):
        return self.name

    def __unicode__(self):
        return self.name


if __name__ == '__main__':
    # doc = docx.Document('2015_Sotrudniki_apparata.docx')
    # doc = docx.Document('2015_Deputaty_(utochnionnye).docx')
    # doc = docx.Document('2017_Deputaty.docx')
    # filename = '2010_Deputaty.docx'
    filename = '2018_Sotrudniki_apparata.docx'
    output_filename = '%s.xml' % filename.split('.')[0].lower()
    doc = docx.Document(filename)

    deputies = []

    if len(doc.tables) > 1:
        table = doc.tables[1]
        row_start = 1
    else:
        table = doc.tables[0]
        row_start = 2

    for i, row in enumerate(table.rows[row_start:]):
        # print('loop')
        data = []
        for idx, cell in enumerate(row.cells):
            if len(row.cells) == 7 and idx == 6:
                data.append(cell.text.encode('utf-8').replace('–', ':').replace('  ', ' ').rstrip())
            elif ('2012' in filename or '2011' in filename) and len(row.cells) == 7:
                if idx == 3 or idx == 4 or idx == 5:
                    data.append(cell.text.encode('utf-8'))
                else:
                    data.append(cell.text.encode('utf-8').replace('\n', ' ').replace('  ', ' ').rstrip())
            elif idx == 9:
                data.append(cell.text.encode('utf-8').replace('–', ':').replace('  ', ' ').rstrip())
            elif idx == 10:
                data.append(cell.text.encode('utf-8').replace('  ', ' ').rstrip())
            else:
                data.append(cell.text.encode('utf-8').replace('\n', ' ').replace('  ', ' ').rstrip())
        # data = [cell.text.encode('utf-8').replace('\n', ' ').replace('  ', ' ').rstrip() for cell in row.cells]
        if len(deputies) > 0 and data[0] == deputies[-1].name or data[0] == '':
            depute = deputies[-1]
        else:
            # if i > 0:
            #     break
            print(data[0])
            depute = Depute()
            depute.set_name(data[0])
            depute.position = data[1]
            depute.set_salary(data[2])
            deputies.append(depute)

        if len(data) == 11:
            depute.set_object_owner(data[3], data[4], data[5])
            depute.set_object_not_owner(data[6], data[7], data[8])
            depute.set_machine(data[9])
        else:
            # if 'аренда' in data[3]:
            #     data[3] = data[3].replace('аренда', '(пользование)')
            # if ('2011' or '2012') in filename:
            if '2012' in filename:
                objects_list = [o.replace(',', '.') for o in data[3].split('\n') if o != '']
                square_list = [s.replace(',', '.') for s in data[4].split('\n') if s != '']
                country_list = [c for c in data[5].split('\n') if c != '']

                for idx in xrange(0, len(objects_list)):
                    square = square_list[idx]
                    country = country_list[idx]
                    depute.write_owner_for_old_files(objects_list[idx], square, country)

            else:
                depute.write_owner_for_old_files(data[3].replace('\n', ''), data[4].replace('\n', ''), data[5].replace('\n', ''))
            depute.set_machine(data[6])

    print('Create xml ...')
    deklarator = DeklaratorXML()

    deklarator.add_info(deputies)

    # with open('2015_sotrudniki_apparata.xml', 'wb') as f:
    # with open('2015_deputaty.xml', 'wb') as f:
    # with open('2015_sotrudniki_apparata.xml', 'wb') as f:
    with open(output_filename, 'wb') as f:
        f.write(deklarator.get_xml())
